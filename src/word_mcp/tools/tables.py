"""Table editing tools for Word documents.

Provides table CRUD operations (create, read, update, add rows/columns) for Word documents.
All functions use zero-based indexing for tables, rows, and columns.

Note: Row deletion (TBL-05) and column deletion (TBL-07) require COM and are not
implemented here. See plan 03-03 for deletion operations.
"""

from pathlib import Path
from typing import Optional, List
from docx.shared import Inches
from ..document_manager import document_manager
from ..logging_config import get_logger

logger = get_logger(__name__)


def create_table(
    path: str,
    rows: int,
    cols: int,
    data: Optional[List[List]] = None,
    style: Optional[str] = None
) -> str:
    """Create a new table in the document.

    The table is appended at the end of the document.

    Args:
        path: Document path or key
        rows: Number of rows (must be > 0)
        cols: Number of columns (must be > 0)
        data: Optional 2D list to populate table. If provided, len(data) must equal
              rows and each inner list must have cols elements.
        style: Optional table style name (e.g., "Table Grid", "Light Shading")

    Returns:
        Success message with table index and count, or error message

    Example:
        create_table(key, 3, 2)  # Creates empty 3x2 table
        create_table(key, 2, 2, data=[["A", "B"], ["C", "D"]])  # Populated table
        create_table(key, 3, 2, style="Table Grid")  # Styled table
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="create_table", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    # Validate dimensions
    if rows <= 0:
        return f"Error: rows must be positive (got {rows})."
    if cols <= 0:
        return f"Error: cols must be positive (got {cols})."

    # Validate data if provided
    if data is not None:
        if len(data) != rows:
            return f"Error: data has {len(data)} rows but table has {rows} rows."
        for i, row_data in enumerate(data):
            if len(row_data) != cols:
                return f"Error: data row {i} has {len(row_data)} items but table has {cols} columns."

    # Create table
    table = doc.add_table(rows=rows, cols=cols)

    # Populate with data if provided
    if data is not None:
        for row_idx in range(rows):
            for col_idx in range(cols):
                table.cell(row_idx, col_idx).text = str(data[row_idx][col_idx])

    # Apply style if provided
    if style is not None:
        try:
            table.style = style
        except KeyError:
            # Get available styles for helpful error message
            available_styles = [s.name for s in doc.styles if s.type == 2]  # type 2 = table styles
            return f"Error: Table style '{style}' not found. Available table styles: {', '.join(available_styles[:10])}..."

    # Get table index (zero-based)
    table_idx = len(doc.tables) - 1
    table_count = len(doc.tables)

    return f"Created table with {rows} rows x {cols} columns (table index: {table_idx}). Document now has {table_count} table(s)."


def list_tables(path: str) -> str:
    """List all tables in the document.

    Shows table index, dimensions, and a preview of the first cell.

    Args:
        path: Document path or key

    Returns:
        Formatted list of tables or error message

    Example output:
        Tables in 'document.docx': 2 table(s)
        [0] 3 rows x 2 cols - first cell: "Header 1"
        [1] 5 rows x 4 cols - first cell: "Name"
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="list_tables", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    tables = doc.tables
    table_count = len(tables)

    # Get filename for display
    filename = Path(path).name if not path.startswith("Untitled-") else path

    if table_count == 0:
        return f"No tables found in '{filename}'."

    # Build header
    lines = [f"Tables in '{filename}': {table_count} table(s)"]

    # List each table
    for idx, table in enumerate(tables):
        row_count = len(table.rows)
        col_count = len(table.columns)

        # Get first cell preview
        try:
            first_cell_text = table.cell(0, 0).text
            # Truncate if longer than 30 chars
            if len(first_cell_text) > 30:
                preview = first_cell_text[:30] + "..."
            else:
                preview = first_cell_text
        except:
            preview = "(empty)"

        lines.append(f"[{idx}] {row_count} rows x {col_count} cols - first cell: \"{preview}\"")

    return "\n".join(lines)


def read_table(
    path: str,
    table_index: int,
    start_row: Optional[int] = None,
    end_row: Optional[int] = None
) -> str:
    """Read the content of a table in formatted grid layout.

    Displays table content with row and column labels. Supports optional row range.

    Args:
        path: Document path or key
        table_index: Zero-based table index
        start_row: Optional starting row (0-based, inclusive)
        end_row: Optional ending row (0-based, inclusive)

    Returns:
        Formatted table grid or error message

    Example output:
        Table 0 (3 rows x 2 cols):
        Row 0: | Header 1 | Header 2 |
        Row 1: | Data A | Data B |
        Row 2: | Data C | Data D |
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="read_table", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    # Validate table index
    table_count = len(doc.tables)
    if table_count == 0:
        return "Error: No tables found in document."
    if table_index < 0 or table_index >= table_count:
        return f"Error: Invalid table_index {table_index}. Document has {table_count} table(s) (valid range: 0-{table_count-1})."

    table = doc.tables[table_index]
    row_count = len(table.rows)
    col_count = len(table.columns)

    # Handle row range
    if start_row is None:
        start_row = 0
    if end_row is None:
        end_row = row_count - 1

    # Validate row range
    if start_row < 0 or start_row >= row_count:
        return f"Error: Invalid start_row {start_row}. Table has {row_count} rows (valid range: 0-{row_count-1})."
    if end_row < 0 or end_row >= row_count:
        return f"Error: Invalid end_row {end_row}. Table has {row_count} rows (valid range: 0-{row_count-1})."
    if start_row > end_row:
        return f"Error: start_row ({start_row}) cannot be greater than end_row ({end_row})."

    # Build header
    lines = [f"Table {table_index} ({row_count} rows x {col_count} cols):"]

    # Build row content
    for row_idx in range(start_row, end_row + 1):
        row_cells = []
        for col_idx in range(col_count):
            try:
                cell_text = table.cell(row_idx, col_idx).text
                # Truncate if longer than 40 chars
                if len(cell_text) > 40:
                    cell_text = cell_text[:40] + "..."
            except:
                # Handle merged cells or inaccessible cells
                cell_text = "?"

            row_cells.append(cell_text)

        # Format row with pipe separators
        row_line = f"Row {row_idx}: | " + " | ".join(row_cells) + " |"
        lines.append(row_line)

    return "\n".join(lines)


def edit_table_cell(path: str, table_index: int, row: int, col: int, text: str) -> str:
    """Edit the content of a specific table cell.

    Note: Using cell.text is acceptable for table cells since table cells don't
    typically have complex run formatting that needs preserving (unlike paragraphs).

    Args:
        path: Document path or key
        table_index: Zero-based table index
        row: Zero-based row index
        col: Zero-based column index
        text: New cell content (converted to string)

    Returns:
        Success message with updated cell preview, or error message

    Example:
        edit_table_cell(key, 0, 1, 2, "Updated text")  # Updates table 0, row 1, col 2
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="edit_table_cell", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    # Validate table index
    table_count = len(doc.tables)
    if table_count == 0:
        return "Error: No tables found in document."
    if table_index < 0 or table_index >= table_count:
        return f"Error: Invalid table_index {table_index}. Document has {table_count} table(s) (valid range: 0-{table_count-1})."

    table = doc.tables[table_index]
    row_count = len(table.rows)
    col_count = len(table.columns)

    # Validate row and column
    if row < 0 or row >= row_count:
        return f"Error: Invalid row {row}. Table has {row_count} rows (valid range: 0-{row_count-1})."
    if col < 0 or col >= col_count:
        return f"Error: Invalid col {col}. Table has {col_count} columns (valid range: 0-{col_count-1})."

    # Update cell content
    table.cell(row, col).text = str(text)

    # Preview: first 50 chars
    text_preview = str(text)[:50] if len(str(text)) <= 50 else str(text)[:50] + "..."

    return f"Updated cell ({row}, {col}) in table {table_index}. New content: '{text_preview}'."


def add_table_row(
    path: str,
    table_index: int,
    data: Optional[List] = None
) -> str:
    """Add a new row to the end of an existing table.

    Args:
        path: Document path or key
        table_index: Zero-based table index
        data: Optional list of cell values. If provided, must have length equal to
              column count. Each value is converted to string.

    Returns:
        Success message with updated table dimensions, or error message

    Example:
        add_table_row(key, 0)  # Adds empty row to table 0
        add_table_row(key, 0, data=["A", "B", "C"])  # Adds populated row
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="add_table_row", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    # Validate table index
    table_count = len(doc.tables)
    if table_count == 0:
        return "Error: No tables found in document."
    if table_index < 0 or table_index >= table_count:
        return f"Error: Invalid table_index {table_index}. Document has {table_count} table(s) (valid range: 0-{table_count-1})."

    table = doc.tables[table_index]
    col_count = len(table.columns)

    # Validate data if provided
    if data is not None:
        if len(data) != col_count:
            return f"Error: data has {len(data)} items but table has {col_count} columns."

    # Add row
    new_row = table.add_row()

    # Populate with data if provided
    if data is not None:
        for col_idx in range(col_count):
            new_row.cells[col_idx].text = str(data[col_idx])

    # Get updated dimensions
    new_row_count = len(table.rows)

    return f"Added row to table {table_index}. Table now has {new_row_count} rows x {col_count} columns."


def add_table_column(
    path: str,
    table_index: int,
    width: Optional[float] = None,
    data: Optional[List] = None
) -> str:
    """Add a new column to the end of an existing table.

    Args:
        path: Document path or key
        table_index: Zero-based table index
        width: Optional column width in inches. Defaults to 1 inch if not provided.
        data: Optional list of cell values. If provided, must have length equal to
              row count. Each value is converted to string.

    Returns:
        Success message with updated table dimensions, or error message

    Example:
        add_table_column(key, 0)  # Adds empty column with default width
        add_table_column(key, 0, width=1.5)  # Adds column with 1.5 inch width
        add_table_column(key, 0, data=["A", "B", "C"])  # Adds populated column
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError as e:
        logger.error("tool_operation_failed", tool="add_table_column", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"

    # Validate table index
    table_count = len(doc.tables)
    if table_count == 0:
        return "Error: No tables found in document."
    if table_index < 0 or table_index >= table_count:
        return f"Error: Invalid table_index {table_index}. Document has {table_count} table(s) (valid range: 0-{table_count-1})."

    table = doc.tables[table_index]
    row_count = len(table.rows)

    # Validate data if provided
    if data is not None:
        if len(data) != row_count:
            return f"Error: data has {len(data)} items but table has {row_count} rows."

    # Add column with specified or default width
    column_width = Inches(width) if width is not None else Inches(1)
    table.add_column(width=column_width)

    # Get new column index (last column)
    new_col_count = len(table.columns)
    new_col_idx = new_col_count - 1

    # Populate with data if provided
    if data is not None:
        for row_idx in range(row_count):
            table.cell(row_idx, new_col_idx).text = str(data[row_idx])

    return f"Added column to table {table_index}. Table now has {row_count} rows x {new_col_count} columns."
