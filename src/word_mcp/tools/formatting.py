"""Text formatting tools for Word documents.

Provides run-level formatting operations (bold, italic, underline, font, size, color).
All formatting operations work at the Run level to preserve existing formatting on other runs.
"""

from typing import Optional
from docx.shared import Pt, RGBColor
from ..document_manager import document_manager
from ..logging_config import get_logger

logger = get_logger(__name__)


def format_text(
    path: str,
    paragraph_index: int,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    font_color: Optional[str] = None,
    run_index: Optional[int] = None
) -> str:
    """Apply formatting to runs in a paragraph.

    CRITICAL: This function works at the Run level to preserve existing formatting.
    It never sets paragraph.text, which would destroy all formatting.

    Args:
        path: Document path or key
        paragraph_index: 0-based paragraph index
        bold: Optional bool (True/False/None). None means "don't change" (preserves python-docx tri-state).
        italic: Optional bool (True/False/None). None means "don't change".
        underline: Optional bool (True/False/None). None means "don't change".
        font_name: Optional str (e.g., "Arial", "Times New Roman"). None means "don't change".
        font_size: Optional float in points (e.g., 12.0 for 12pt). None means "don't change".
        font_color: Optional str in hex format "#RRGGBB" (e.g., "#FF0000" for red). None means "don't change".
        run_index: Optional int (0-based). If None, applies to ALL runs in paragraph. If specified, applies only to that run.

    Returns:
        Success message listing what was changed, or error message

    Examples:
        format_text(key, 0, bold=True)  # Bold all runs in paragraph 0
        format_text(key, 1, font_name="Arial", font_size=12.0, run_index=0)  # Format only run 0 in paragraph 1
        format_text(key, 2, font_color="#FF0000")  # Red text for all runs in paragraph 2
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError:
        return f"Error: No document open at '{path}'. Use open_document or create_document first."

    para_count = len(doc.paragraphs)

    # Validate paragraph index
    if paragraph_index < 0 or paragraph_index >= para_count:
        return f"Error: Invalid paragraph index {paragraph_index}. Document has {para_count} paragraphs (valid range: 0-{para_count-1})."

    para = doc.paragraphs[paragraph_index]

    # Check if paragraph has runs
    if len(para.runs) == 0:
        return f"Error: Paragraph {paragraph_index} has no runs (empty paragraph). Cannot apply formatting."

    # Validate font_color format if provided
    if font_color is not None:
        if not (len(font_color) == 7 and font_color[0] == "#"):
            return f"Error: Invalid font_color format '{font_color}'. Expected '#RRGGBB' (e.g., '#FF0000')."
        try:
            # Validate hex chars
            int(font_color[1:3], 16)
            int(font_color[3:5], 16)
            int(font_color[5:7], 16)
        except ValueError:
            return f"Error: Invalid font_color format '{font_color}'. Expected '#RRGGBB' with valid hex digits."

    # Determine which runs to format
    if run_index is not None:
        # Validate run_index
        if run_index < 0 or run_index >= len(para.runs):
            return f"Error: Invalid run_index {run_index}. Paragraph {paragraph_index} has {len(para.runs)} runs (valid range: 0-{len(para.runs)-1})."
        runs_to_format = [para.runs[run_index]]
        target_desc = f"run {run_index}"
    else:
        runs_to_format = para.runs
        target_desc = f"all {len(para.runs)} runs"

    # Apply formatting to each target run
    changes = []
    for run in runs_to_format:
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if font_name is not None:
            run.font.name = font_name
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_color is not None:
            # Parse hex color
            r = int(font_color[1:3], 16)
            g = int(font_color[3:5], 16)
            b = int(font_color[5:7], 16)
            run.font.color.rgb = RGBColor(r, g, b)

    # Build change description
    if bold is not None:
        changes.append(f"bold={bold}")
    if italic is not None:
        changes.append(f"italic={italic}")
    if underline is not None:
        changes.append(f"underline={underline}")
    if font_name is not None:
        changes.append(f"font='{font_name}'")
    if font_size is not None:
        changes.append(f"size={font_size}pt")
    if font_color is not None:
        changes.append(f"color={font_color}")

    if not changes:
        return f"No formatting changes specified for paragraph {paragraph_index}."

    changes_str = ", ".join(changes)
    return f"Applied formatting to paragraph {paragraph_index} ({target_desc}): {changes_str}"


def get_paragraph_formatting(path: str, paragraph_index: int) -> str:
    """Get formatting details for all runs in a paragraph.

    Returns formatting information for each run: index, text preview, bold, italic,
    underline, font name, font size, and font color.

    Args:
        path: Document path or key
        paragraph_index: 0-based paragraph index

    Returns:
        Formatted string with per-run formatting details, or error message

    Example output:
        Paragraph 0 formatting (3 runs):
        [0] "Hello worl..." bold=True italic=False underline=False font=Arial size=12pt color=#FF0000
        [1] " world" bold=False italic=True underline=False font=Arial size=12pt color=inherited
        [2] "!" bold=False italic=False underline=False font=Times New Roman size=14pt color=#0000FF
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError:
        return f"Error: No document open at '{path}'. Use open_document or create_document first."

    para_count = len(doc.paragraphs)

    # Validate paragraph index
    if paragraph_index < 0 or paragraph_index >= para_count:
        return f"Error: Invalid paragraph index {paragraph_index}. Document has {para_count} paragraphs (valid range: 0-{para_count-1})."

    para = doc.paragraphs[paragraph_index]

    if len(para.runs) == 0:
        return f"Paragraph {paragraph_index} has no runs (empty paragraph)."

    # Build header
    lines = [f"Paragraph {paragraph_index} formatting ({len(para.runs)} runs):"]

    # Build per-run details
    for i, run in enumerate(para.runs):
        # Text preview: first 30 chars
        text = run.text
        if len(text) > 30:
            text_preview = f'"{text[:30]}..."'
        else:
            text_preview = f'"{text}"'

        # Formatting properties (show "inherited" for None values)
        bold_val = run.bold if run.bold is not None else "inherited"
        italic_val = run.italic if run.italic is not None else "inherited"
        underline_val = run.underline if run.underline is not None else "inherited"

        font_name = run.font.name if run.font.name is not None else "inherited"

        if run.font.size is not None:
            # Convert EMUs to points
            font_size = f"{run.font.size.pt}pt"
        else:
            font_size = "inherited"

        if run.font.color.rgb is not None:
            # Convert RGBColor to hex string
            rgb = run.font.color.rgb
            font_color = f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
        else:
            font_color = "inherited"

        lines.append(
            f"[{i}] {text_preview} bold={bold_val} italic={italic_val} underline={underline_val} "
            f"font={font_name} size={font_size} color={font_color}"
        )

    return "\n".join(lines)
