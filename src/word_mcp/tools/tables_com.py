"""
COM-based table editing tools for word-mcp.

This module provides table row and column deletion, and tracked table cell editing,
using win32com automation. python-docx does not support row/column deletion or
tracked cell edits, so COM automation is required.

All functions use the bridge pattern:
1. Validate document open in DocumentManager
2. Validate document saved to disk (COM needs file on disk)
3. Open via COM (WordApplication context manager)
4. Perform COM-based operation
5. Save and close via COM
6. Reload python-docx document to sync state
"""

from pathlib import Path
from docx import Document
from ..document_manager import document_manager
from ..com_pool import com_pool
from ..logging_config import get_logger

logger = get_logger(__name__)


def delete_table_row(path: str, table_index: int, row_index: int) -> str:
    """
    Delete a row from an existing table using COM automation (TBL-05).

    python-docx does not support row deletion, so this function uses win32com
    to delete a row and maintain synchronization with the in-memory Document.

    PREREQUISITE: Document must be saved to disk first (COM opens files from disk).

    Args:
        path: Path or key of open document
        table_index: Zero-based table index in the document
        row_index: Zero-based row index within the table

    Returns:
        Success message with updated row count, or error message prefixed with "Error:"

    Examples:
        >>> delete_table_row("C:/Documents/report.docx", 0, 2)
        "Deleted row 2 from table 0. Table now has 4 rows."

        >>> delete_table_row("Untitled-1", 0, 0)
        "Error: Document must be saved to disk before COM operations. Use save_document_as first."

    Design notes:
        - Requires COM automation: Document must be saved to disk
        - Bridge pattern: Uses COM to delete row, then reloads python-docx
        - Zero-based indexing: Converts to 1-based for COM internally
        - Index validation: Checks table and row exist before deleting
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before COM operations. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before COM operations. Use save_document first."

        # Validate table_index is within bounds (using python-docx for validation)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Error: Invalid table index {table_index}. Document has {len(doc.tables)} table(s) (valid range: 0-{len(doc.tables) - 1})."

        # Use COM to delete row
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Convert 0-based to 1-based for COM
                com_table_index = table_index + 1
                com_row_index = row_index + 1

                # Validate table exists in COM document
                if com_table_index > com_doc.Tables.Count:
                    return f"Error: Invalid table index {table_index}. Document has {com_doc.Tables.Count} table(s)."

                # Get table
                table = com_doc.Tables(com_table_index)

                # Validate row exists
                if com_row_index < 1 or com_row_index > table.Rows.Count:
                    return f"Error: Invalid row index {row_index}. Table {table_index} has {table.Rows.Count} row(s) (valid range: 0-{table.Rows.Count - 1})."

                # Delete row
                table.Rows(com_row_index).Delete()

                # Get updated row count
                updated_row_count = table.Rows.Count

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        return f"Deleted row {row_index} from table {table_index}. Table now has {updated_row_count} rows."

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"


def delete_table_column(path: str, table_index: int, col_index: int) -> str:
    """
    Delete a column from an existing table using COM automation (TBL-07).

    python-docx does not support column deletion, so this function uses win32com
    to delete a column and maintain synchronization with the in-memory Document.

    PREREQUISITE: Document must be saved to disk first (COM opens files from disk).

    Args:
        path: Path or key of open document
        table_index: Zero-based table index in the document
        col_index: Zero-based column index within the table

    Returns:
        Success message with updated column count, or error message prefixed with "Error:"

    Examples:
        >>> delete_table_column("C:/Documents/report.docx", 0, 1)
        "Deleted column 1 from table 0. Table now has 3 columns."

        >>> delete_table_column("Untitled-1", 0, 0)
        "Error: Document must be saved to disk before COM operations. Use save_document_as first."

    Design notes:
        - Requires COM automation: Document must be saved to disk
        - Bridge pattern: Uses COM to delete column, then reloads python-docx
        - Zero-based indexing: Converts to 1-based for COM internally
        - Index validation: Checks table and column exist before deleting
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before COM operations. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before COM operations. Use save_document first."

        # Validate table_index is within bounds (using python-docx for validation)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Error: Invalid table index {table_index}. Document has {len(doc.tables)} table(s) (valid range: 0-{len(doc.tables) - 1})."

        # Use COM to delete column
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Convert 0-based to 1-based for COM
                com_table_index = table_index + 1
                com_col_index = col_index + 1

                # Validate table exists in COM document
                if com_table_index > com_doc.Tables.Count:
                    return f"Error: Invalid table index {table_index}. Document has {com_doc.Tables.Count} table(s)."

                # Get table
                table = com_doc.Tables(com_table_index)

                # Validate column exists
                if com_col_index < 1 or com_col_index > table.Columns.Count:
                    return f"Error: Invalid column index {col_index}. Table {table_index} has {table.Columns.Count} column(s) (valid range: 0-{table.Columns.Count - 1})."

                # Delete column
                table.Columns(com_col_index).Delete()

                # Get updated column count
                updated_col_count = table.Columns.Count

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        return f"Deleted column {col_index} from table {table_index}. Table now has {updated_col_count} columns."

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"


def tracked_edit_table_cell(
    path: str, table_index: int, row_index: int, col_index: int,
    new_text: str, author: str = "Claude"
) -> str:
    """
    Edit a table cell creating tracked Deletion + Insertion revisions (Phase 6).

    This uses COM automation to replace a table cell's text so Word records it as
    tracked changes (old text = Deletion, new text = Insertion). The document MUST
    have tracking enabled first (via enable_tracked_changes).

    Cell content replacement works by getting the cell range, trimming the trailing
    cell end marker (\\r\\x07), and assigning new text to the range. This creates
    a tracked Deletion of old content and tracked Insertion of new content, visible
    in Word as strikethrough + colored text.

    PREREQUISITE: Document must be saved to disk first (COM opens files from disk).
    PREREQUISITE: Tracked changes must be enabled (call enable_tracked_changes first).

    Args:
        path: Path or key of open document
        table_index: Zero-based table index in the document
        row_index: Zero-based row index within the table
        col_index: Zero-based column index within the table
        new_text: New text content to replace existing cell text
        author: Author name for the tracked changes (default: "Claude")

    Returns:
        Success message with before/after preview, or error message prefixed with "Error:"

    Examples:
        >>> tracked_edit_table_cell("C:/Documents/report.docx", 0, 1, 0, "Updated value", "Claude")
        "Edited tracked table 0, cell (1, 0). Was: 'Old value' -> Now: 'Updated value'. Changes tracked as revisions by 'Claude'."

        >>> tracked_edit_table_cell("Untitled-1", 0, 0, 0, "text")
        "Error: Document must be saved to disk before tracked editing. Use save_document_as first."

        >>> tracked_edit_table_cell("C:/Documents/no-track.docx", 0, 0, 0, "text")
        "Error: Tracked changes are not enabled on this document. Call enable_tracked_changes first."

    Design notes:
        - Requires COM automation: Document must be saved to disk
        - Requires tracked changes enabled: Returns error if TrackRevisions=False
        - Bridge pattern: Uses COM for tracked edit, then reloads python-docx
        - Zero-based indexing: All indexes are 0-based, converted to 1-based for COM
        - Cell end marker: COM cell ranges end with \\r\\x07; Range.End is adjusted to exclude it
        - Author attribution: Sets UserName in Word before editing
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before tracked editing. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before tracked editing. Use save_document first."

        # Validate table_index is within bounds (using python-docx for validation)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Error: Invalid table index {table_index}. Document has {len(doc.tables)} table(s) (valid range: 0-{len(doc.tables) - 1})."

        # Use COM to edit table cell with tracked changes
        old_text = ""
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Verify tracking is enabled
                if not com_doc.TrackRevisions:
                    return "Error: Tracked changes are not enabled on this document. Call enable_tracked_changes first."

                # Set author for new revisions
                word.UserName = author

                # Convert 0-based to 1-based for COM
                com_table_index = table_index + 1
                com_row_index = row_index + 1
                com_col_index = col_index + 1

                # Validate table exists in COM document
                if com_table_index > com_doc.Tables.Count:
                    return f"Error: Invalid table index {table_index}. Document has {com_doc.Tables.Count} table(s)."

                # Get table
                table = com_doc.Tables(com_table_index)

                # Validate row exists
                if com_row_index < 1 or com_row_index > table.Rows.Count:
                    return f"Error: Invalid row index {row_index}. Table {table_index} has {table.Rows.Count} row(s) (valid range: 0-{table.Rows.Count - 1})."

                # Validate column exists
                if com_col_index < 1 or com_col_index > table.Columns.Count:
                    return f"Error: Invalid column index {col_index}. Table {table_index} has {table.Columns.Count} column(s) (valid range: 0-{table.Columns.Count - 1})."

                # Get the cell
                cell = table.Cell(com_row_index, com_col_index)

                # Capture old text for confirmation
                old_text = cell.Range.Text

                # Get cell range and trim trailing cell end marker (\r\x07)
                # Cell text always ends with \r\x07 (paragraph mark + cell end marker)
                # We must exclude this from the range before setting new text
                cell_range = cell.Range
                # Trim trailing markers: strip \r\x07 (2 chars) from the end of the range
                cell_range.End = cell_range.End - 1  # Exclude cell end marker (\x07)
                if cell_range.Text.endswith('\r'):
                    cell_range.End = cell_range.End - 1  # Exclude paragraph mark (\r)

                # Replace text (creates Deletion + Insertion revisions when tracking is on)
                cell_range.Text = new_text

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        # Prepare success message with text previews
        old_preview_text = old_text.replace('\r', ' ').replace('\x07', '').strip()
        old_preview = old_preview_text[:50] + "..." if len(old_preview_text) > 50 else old_preview_text
        new_preview = new_text[:50] + "..." if len(new_text) > 50 else new_text
        return (
            f"Edited tracked table {table_index}, cell ({row_index}, {col_index}). "
            f"Was: '{old_preview}' -> Now: '{new_preview}'. "
            f"Changes tracked as revisions by '{author}'."
        )

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"
