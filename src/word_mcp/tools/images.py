"""Image insertion and manipulation tools for Word documents.

Provides inline image insertion, resizing, and listing capabilities using python-docx.
All images are inserted as InlineShapes (inline with text, not floating).
"""

from pathlib import Path
from docx.shared import Inches
from ..document_manager import document_manager
from ..logging_config import get_logger

logger = get_logger(__name__)


def insert_image(
    path: str,
    image_path: str,
    width: float = None,
    height: float = None,
    paragraph_index: int = None
) -> str:
    """Insert an image into the document as an InlineShape.

    By default, appends a new paragraph with the image at the end of the document.
    Can optionally insert into an existing paragraph.

    Args:
        path: Document path or key
        image_path: Path to image file (png, jpg, gif, bmp, etc.)
        width: Optional width in inches. If only width specified, aspect ratio preserved.
        height: Optional height in inches. If only height specified, aspect ratio preserved.
        paragraph_index: Optional 0-based paragraph index. If None, creates new paragraph at end.

    Returns:
        Success message with image details and count, or error message

    Examples:
        insert_image(key, "logo.png")  # Insert at end, original size
        insert_image(key, "photo.jpg", width=4.0)  # Insert at end, 4 inches wide, aspect preserved
        insert_image(key, "icon.png", width=1.0, height=1.0, paragraph_index=2)  # Insert in paragraph 2, 1x1 inches

    NOTE: If both width and height specified, image may be distorted if aspect ratio doesn't match.
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError:
        return f"Error: No document open at '{path}'. Use open_document or create_document first."

    # Validate image file exists
    img_path = Path(image_path)
    if not img_path.exists():
        return f"Error: Image file not found: {image_path}"

    # Prepare width/height arguments for python-docx
    width_arg = Inches(width) if width is not None else None
    height_arg = Inches(height) if height is not None else None

    # Insert image
    if paragraph_index is None:
        # Append new paragraph with image at end of document
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(str(img_path), width=width_arg, height=height_arg)
        insert_location = len(doc.paragraphs) - 1
    else:
        # Insert in existing paragraph
        para_count = len(doc.paragraphs)
        if paragraph_index < 0 or paragraph_index >= para_count:
            return f"Error: Invalid paragraph index {paragraph_index}. Document has {para_count} paragraphs (valid range: 0-{para_count-1})."

        para = doc.paragraphs[paragraph_index]
        run = para.add_run()
        run.add_picture(str(img_path), width=width_arg, height=height_arg)
        insert_location = paragraph_index

    # Get total inline images count
    total_images = len(doc.inline_shapes)

    # Get filename for display
    image_filename = img_path.name

    # Build dimensions string
    if width is not None and height is not None:
        dims = f"width={width}in, height={height}in"
    elif width is not None:
        dims = f"width={width}in (aspect preserved)"
    elif height is not None:
        dims = f"height={height}in (aspect preserved)"
    else:
        dims = "original size"

    return f"Inserted image '{image_filename}' ({dims}) at paragraph {insert_location}. Document has {total_images} inline images."


def resize_image(
    path: str,
    image_index: int,
    width: float = None,
    height: float = None
) -> str:
    """Resize an existing inline image by its index.

    NOTE: When resizing an existing image, python-docx does NOT automatically preserve
    aspect ratio. If you set only width, the height remains unchanged (and vice versa).
    To preserve aspect ratio, you must calculate and set both dimensions.

    Args:
        path: Document path or key
        image_index: 0-based index into doc.inline_shapes collection
        width: Optional width in inches
        height: Optional height in inches

    Returns:
        Success message with new dimensions, or error message

    Examples:
        resize_image(key, 0, width=5.0)  # Set width to 5 inches (height unchanged)
        resize_image(key, 1, width=3.0, height=2.0)  # Set both dimensions
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError:
        return f"Error: No document open at '{path}'. Use open_document or create_document first."

    # Validate at least one dimension provided
    if width is None and height is None:
        return "Error: At least one dimension (width or height) must be provided."

    # Validate image_index
    image_count = len(doc.inline_shapes)
    if image_count == 0:
        return "Error: Document has no inline images."

    if image_index < 0 or image_index >= image_count:
        return f"Error: Invalid image_index {image_index}. Document has {image_count} inline images (valid range: 0-{image_count-1})."

    # Get the inline shape
    shape = doc.inline_shapes[image_index]

    # Apply new dimensions
    if width is not None:
        shape.width = Inches(width)
    if height is not None:
        shape.height = Inches(height)

    # Get current dimensions for confirmation
    current_width = shape.width.inches
    current_height = shape.height.inches

    return f"Resized image {image_index} to width={current_width:.2f}in, height={current_height:.2f}in."


def list_images(path: str) -> str:
    """List all inline images in the document with dimensions.

    Returns index, width, and height for each inline image in the document.

    Args:
        path: Document path or key

    Returns:
        Formatted string with image list and dimensions, or error message

    Example output:
        Inline images in 'document.docx': 3 image(s)
        [0] 4.50in x 3.00in
        [1] 2.00in x 2.00in
        [2] 6.00in x 4.50in
    """
    try:
        doc = document_manager.get_document(path)
    except ValueError:
        return f"Error: No document open at '{path}'. Use open_document or create_document first."

    # Get filename for display
    if path.startswith("Untitled-"):
        filename = path
    else:
        filename = Path(path).name

    # Get inline images
    images = doc.inline_shapes
    image_count = len(images)

    if image_count == 0:
        return f"No inline images found in '{filename}'."

    # Build header
    lines = [f"Inline images in '{filename}': {image_count} image(s)"]

    # Build image list
    for i, shape in enumerate(images):
        # Get dimensions in inches
        width_inches = shape.width.inches
        height_inches = shape.height.inches

        lines.append(f"[{i}] {width_inches:.2f}in x {height_inches:.2f}in")

    return "\n".join(lines)
