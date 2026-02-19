"""
Tracked editing tools for word-mcp.

This module provides COM-based paragraph editing functions that create tracked
changes when tracking is enabled. These are ADDITIONAL tools (not replacements)
for Phase 1's python-docx tools.

User chooses:
- Phase 1 tools (add_paragraph, edit_paragraph, delete_paragraph) for fast untracked edits
- Phase 2 tools (tracked_add_paragraph, tracked_edit_paragraph, tracked_delete_paragraph) for tracked edits

All functions use the bridge pattern:
1. Validate document open in DocumentManager
2. Validate document saved to disk (COM needs file on disk)
3. Open via COM (WordApplication context manager)
4. Verify tracking enabled (TrackRevisions == True)
5. Set UserName to author parameter
6. Perform COM-based edit
7. Save and close via COM
8. Reload python-docx document to sync state

Phase 6 addition: _translate_paragraph_index translates python-docx body paragraph
indexes to COM paragraph indexes, skipping table cell paragraphs. COM's
Document.Paragraphs includes paragraphs inside table cells; python-docx's
document.paragraphs does not.
"""

from pathlib import Path
from docx import Document
from ..document_manager import document_manager
from ..com_pool import com_pool
from ..logging_config import get_logger

logger = get_logger(__name__)

# wdWithInTable constant value for COM Information() call
_WD_WITH_IN_TABLE = 12


def _translate_paragraph_index(com_doc, python_docx_index: int) -> int:
    """
    Translate a python-docx body paragraph index to the correct COM paragraph index.

    COM's Document.Paragraphs includes paragraphs inside table cells while
    python-docx's document.paragraphs excludes them. This function iterates the
    COM paragraph collection, counts only body-level paragraphs (those NOT inside
    a table cell), and returns the COM 1-based index corresponding to the
    python-docx 0-based index.

    Args:
        com_doc: Open COM Document object
        python_docx_index: Zero-based body paragraph index (as used by python-docx)

    Returns:
        COM 1-based paragraph index for the same body paragraph

    Raises:
        IndexError: If python_docx_index is out of range for the body paragraph count
        ValueError: If python_docx_index is negative
    """
    if python_docx_index < 0:
        raise ValueError(f"Paragraph index must be non-negative, got {python_docx_index}")

    total_com_paragraphs = com_doc.Paragraphs.Count
    body_paragraph_count = 0

    for com_index in range(1, total_com_paragraphs + 1):
        para = com_doc.Paragraphs(com_index)

        # Check if this paragraph is inside a table cell.
        # wdWithInTable (constant value 12) returns True for paragraphs inside tables.
        try:
            in_table = para.Range.Information(_WD_WITH_IN_TABLE)
        except Exception:
            # If Information() fails for any reason, treat as body paragraph
            in_table = False

        if not in_table:
            # This is a body-level paragraph
            if body_paragraph_count == python_docx_index:
                return com_index
            body_paragraph_count += 1

    # If we reach here, the index is out of range
    raise IndexError(
        f"Paragraph index {python_docx_index} is out of range. "
        f"Document has {body_paragraph_count} body paragraph(s) "
        f"(valid range: 0-{body_paragraph_count - 1})."
    )


def tracked_add_paragraph(
    path: str, text: str, position: str = "end", author: str = "Claude",
    expected_text: str = None
) -> str:
    """
    Add a paragraph that creates an Insertion revision (TRACK-03).

    This uses COM automation to add a paragraph that Word records as a tracked
    change. The document MUST have tracking enabled first (via enable_tracked_changes).

    INDEX TRANSLATION: Uses _translate_paragraph_index to map python-docx body
    paragraph indexes to COM paragraph indexes, so that documents with tables are
    handled correctly.

    Args:
        path: Path or key of open document
        text: Text content for the new paragraph
        position: "end" to append, or zero-based index string to insert before
        author: Author name for the tracked change (default: "Claude")
        expected_text: Optional content verification string. If provided and position
                       is not "end", the paragraph at that position must contain this
                       text (case-sensitive partial match). If it does not match, the
                       insert is refused with an error message. This prevents inserting
                       at the wrong location when paragraph indexes have shifted.

    Returns:
        Success message or error message prefixed with "Error:"

    Examples:
        >>> tracked_add_paragraph("C:/Documents/report.docx", "New paragraph", "end", "Claude")
        "Added tracked paragraph at end: 'New paragraph'. Revision will appear as insertion by 'Claude'."

        >>> tracked_add_paragraph("C:/Documents/report.docx", "First", "0", "Claude")
        "Added tracked paragraph at 0: 'First'. Revision will appear as insertion by 'Claude'."

        >>> tracked_add_paragraph("C:/Documents/report.docx", "New intro", "5", "Claude", expected_text="Introduction")
        "Error: Content verification failed for paragraph 5. Expected text containing 'Introduction' but found: 'Background section content here'. The paragraph may have shifted -- re-read the document."
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before tracked editing. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before tracked editing. Use save_document first."

        # Validate position parameter
        if position != "end":
            try:
                position_int = int(position)
                if position_int < 0:
                    return f"Error: Invalid position {position}. Must be 'end' or a non-negative integer."
            except ValueError:
                return f"Error: Invalid position '{position}'. Must be 'end' or a zero-based integer index."

        # Use COM to add paragraph with tracked changes
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Verify tracking is enabled
                if not com_doc.TrackRevisions:
                    return "Error: Tracked changes are not enabled on this document. Call enable_tracked_changes first."

                # Set author for new revisions
                word.UserName = author

                # Add paragraph
                if position == "end":
                    # Append to end: InsertAfter with \r creates new paragraph
                    com_doc.Content.InsertAfter("\r" + text)
                else:
                    position_int = int(position)

                    # Translate python-docx index to COM index
                    try:
                        com_index = _translate_paragraph_index(com_doc, position_int)
                    except (IndexError, ValueError) as e:
                        return f"Error: {str(e)}"

                    # Content verification before inserting
                    if expected_text is not None:
                        para_text = com_doc.Paragraphs(com_index).Range.Text
                        if expected_text not in para_text:
                            actual_preview = para_text[:80].replace('\r', ' ').replace('\x07', '')
                            return (
                                f"Error: Content verification failed for paragraph {position_int}. "
                                f"Expected text containing '{expected_text}' but found: "
                                f"'{actual_preview}'. The paragraph may have shifted -- re-read the document."
                            )

                    # Get paragraph range and insert before
                    para_range = com_doc.Paragraphs(com_index).Range
                    para_range.InsertBefore(text + "\r")

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        # Prepare success message
        text_preview = text[:50] + "..." if len(text) > 50 else text
        return f"Added tracked paragraph at {position}: '{text_preview}'. Revision will appear as insertion by '{author}'."

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"


def tracked_edit_paragraph(
    path: str, index: int, new_text: str, author: str = "Claude",
    expected_text: str = None
) -> str:
    """
    Edit paragraph text creating Deletion + Insertion revisions (TRACK-03).

    This uses COM automation to replace paragraph text so Word records it as
    tracked changes (old text = Deletion, new text = Insertion). The document
    MUST have tracking enabled first (via enable_tracked_changes).

    INDEX TRANSLATION: Uses _translate_paragraph_index to map python-docx body
    paragraph indexes to COM paragraph indexes, so that documents with tables are
    handled correctly.

    Args:
        path: Path or key of open document
        index: Zero-based paragraph index to edit (python-docx body paragraph index)
        new_text: New text content to replace existing text
        author: Author name for the tracked changes (default: "Claude")
        expected_text: Optional content verification string. If provided, the target
                       paragraph must contain this text (case-sensitive partial match)
                       before the edit proceeds. If it does not match, the edit is
                       refused with a descriptive error message. Use this to guard
                       against editing the wrong paragraph when indexes may have shifted.

    Returns:
        Success message or error message prefixed with "Error:"

    Examples:
        >>> tracked_edit_paragraph("C:/Documents/report.docx", 0, "Updated text", "Claude")
        "Edited tracked paragraph 0. Was: 'Original text' -> Now: 'Updated text'. Changes tracked as revisions by 'Claude'."

        >>> tracked_edit_paragraph("C:/Documents/report.docx", 40, "New section text", "Claude", expected_text="Section 4")
        "Edited tracked paragraph 40. Was: 'Section 4 of Task Order...' -> Now: 'New section text'. Changes tracked as revisions by 'Claude'."

        >>> tracked_edit_paragraph("C:/Documents/report.docx", 40, "New text", "Claude", expected_text="Section 4")
        "Error: Content verification failed for paragraph 40. Expected text containing 'Section 4' but found: 'DELIVERABLE'. The paragraph may have shifted -- re-read the document."
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before tracked editing. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before tracked editing. Use save_document first."

        # Use COM to edit paragraph with tracked changes
        old_text = ""
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Verify tracking is enabled
                if not com_doc.TrackRevisions:
                    return "Error: Tracked changes are not enabled on this document. Call enable_tracked_changes first."

                # Set author for new revisions
                word.UserName = author

                # Translate python-docx index to COM index
                try:
                    com_index = _translate_paragraph_index(com_doc, index)
                except (IndexError, ValueError) as e:
                    return f"Error: {str(e)}"

                # Get paragraph range
                para_range = com_doc.Paragraphs(com_index).Range

                # Capture old text for confirmation
                old_text = para_range.Text

                # Content verification before editing
                if expected_text is not None:
                    if expected_text not in old_text:
                        actual_preview = old_text[:80].replace('\r', ' ').replace('\x07', '')
                        return (
                            f"Error: Content verification failed for paragraph {index}. "
                            f"Expected text containing '{expected_text}' but found: "
                            f"'{actual_preview}'. The paragraph may have shifted -- re-read the document."
                        )

                # Strip trailing paragraph mark from range for replacement
                # Range.Text includes \r at the end, we need to exclude it for clean replacement
                if para_range.Text.endswith('\r'):
                    para_range.End = para_range.End - 1

                # Replace text (creates Deletion + Insertion revisions when tracking is on)
                para_range.Text = new_text

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        # Prepare success message with text previews
        old_preview = old_text[:50].strip() + "..." if len(old_text) > 50 else old_text.strip()
        new_preview = new_text[:50] + "..." if len(new_text) > 50 else new_text
        return f"Edited tracked paragraph {index}. Was: '{old_preview}' -> Now: '{new_preview}'. Changes tracked as revisions by '{author}'."

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"


def tracked_delete_paragraph(
    path: str, index: int, author: str = "Claude", expected_text: str = None
) -> str:
    """
    Delete a paragraph creating a Deletion revision (TRACK-03).

    This uses COM automation to delete a paragraph so Word records it as a
    tracked change (strikethrough in Word). The document MUST have tracking
    enabled first (via enable_tracked_changes).

    INDEX TRANSLATION: Uses _translate_paragraph_index to map python-docx body
    paragraph indexes to COM paragraph indexes, so that documents with tables are
    handled correctly.

    INDEX SHIFT WARNING: After deletion, remaining paragraphs shift down.
    Re-read the document before performing additional operations.

    Args:
        path: Path or key of open document
        index: Zero-based paragraph index to delete (python-docx body paragraph index)
        author: Author name for the tracked change (default: "Claude")
        expected_text: Optional content verification string. If provided, the target
                       paragraph must contain this text (case-sensitive partial match)
                       before the delete proceeds. If it does not match, the delete is
                       refused with a descriptive error message. Use this to guard
                       against deleting the wrong paragraph when indexes may have shifted.

    Returns:
        Success message or error message prefixed with "Error:"

    Examples:
        >>> tracked_delete_paragraph("C:/Documents/report.docx", 2, "Claude")
        "Deleted tracked paragraph 2 ('Old text'). Deletion tracked as revision by 'Claude'. Remaining paragraphs have shifted -- re-read document to get updated indexes."

        >>> tracked_delete_paragraph("C:/Documents/report.docx", 5, "Claude", expected_text="Obsolete section")
        "Error: Content verification failed for paragraph 5. Expected text containing 'Obsolete section' but found: 'Introduction text here'. The paragraph may have shifted -- re-read the document."
    """
    try:
        # Validate document is open in DocumentManager
        key = path if path.startswith("Untitled-") else str(Path(path).resolve())
        doc = document_manager.get_document(key)

        # Check file exists on disk (COM requires saved file)
        if key.startswith("Untitled-"):
            return "Error: Document must be saved to disk before tracked editing. Use save_document_as first."

        if not Path(key).exists():
            return "Error: Document must be saved to disk before tracked editing. Use save_document first."

        # Use COM to delete paragraph with tracked changes
        deleted_text = ""
        try:
            with com_pool.get_word_app() as word:
                com_doc = word.Documents.Open(key)

                # Verify tracking is enabled
                if not com_doc.TrackRevisions:
                    return "Error: Tracked changes are not enabled on this document. Call enable_tracked_changes first."

                # Set author for new revisions
                word.UserName = author

                # Translate python-docx index to COM index
                try:
                    com_index = _translate_paragraph_index(com_doc, index)
                except (IndexError, ValueError) as e:
                    return f"Error: {str(e)}"

                # Get paragraph range
                para_range = com_doc.Paragraphs(com_index).Range

                # Capture text for confirmation
                deleted_text = para_range.Text

                # Content verification before deleting
                if expected_text is not None:
                    if expected_text not in deleted_text:
                        actual_preview = deleted_text[:80].replace('\r', ' ').replace('\x07', '')
                        return (
                            f"Error: Content verification failed for paragraph {index}. "
                            f"Expected text containing '{expected_text}' but found: "
                            f"'{actual_preview}'. The paragraph may have shifted -- re-read the document."
                        )

                # Delete (creates Deletion revision when tracking is on)
                para_range.Delete()

                # Save and close
                com_doc.Save()
                com_doc.Close()

        except Exception as e:
            logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
            return f"Error: COM automation failed: {str(e)}. Ensure Microsoft Word is installed."

        # Reload python-docx document to sync in-memory state
        document_manager._documents[key] = Document(key)

        # Prepare success message with text preview and index shift warning
        text_preview = deleted_text[:50].strip() + "..." if len(deleted_text) > 50 else deleted_text.strip()
        return f"Deleted tracked paragraph {index} ('{text_preview}'). Deletion tracked as revision by '{author}'. Remaining paragraphs have shifted -- re-read document to get updated indexes."

    except ValueError:
        return f"Error: Document not open: {path}"
    except Exception as e:
        logger.error("tool_operation_failed", tool="unknown", error=str(e), error_type=type(e).__name__)
        return f"Error: {str(e)}"
