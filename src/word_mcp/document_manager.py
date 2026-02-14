"""
Document lifecycle management for word-mcp.

This module provides the DocumentManager class, which maintains in-memory state
for all open Word documents. Documents are keyed by their absolute file paths
(or temporary "Untitled-N" keys for unsaved documents).

Key behaviors:
- Multi-document support: multiple documents can be open simultaneously
- Error-on-overwrite: create_document() raises FileExistsError if target path exists
- Explicit save only: documents are NOT auto-saved; changes persist only via save_document()
- Path normalization: all paths converted to absolute using Path.resolve()
"""

import os
from pathlib import Path
from typing import Dict, Optional

from docx import Document


class DocumentManager:
    """
    Manages in-memory state for multiple open Word documents.

    Documents are stored in a dictionary keyed by their absolute file path
    (or temporary "Untitled-N" key for unsaved documents). All path operations
    normalize to absolute paths to avoid key collisions.
    """

    def __init__(self):
        self._documents: Dict[str, Document] = {}
        self._untitled_counter: int = 0

    def _next_untitled(self) -> str:
        """
        Generate next available "Untitled-N" key.

        Returns:
            Unique untitled key like "Untitled-1", "Untitled-2", etc.
        """
        self._untitled_counter += 1
        return f"Untitled-{self._untitled_counter}"

    def _normalize_path(self, path: str) -> str:
        """
        Convert path to absolute, normalized form.

        Args:
            path: File path (absolute or relative)

        Returns:
            Absolute path as string
        """
        return str(Path(path).resolve())

    def create_document(self, path: Optional[str] = None) -> tuple[str, Document]:
        """
        Create a new blank Word document in memory.

        Args:
            path: Optional file path for the document. If provided and file exists,
                  raises FileExistsError (error-on-overwrite behavior). If None,
                  generates temporary "Untitled-N" key.

        Returns:
            Tuple of (key, Document) where key is the path or "Untitled-N"

        Raises:
            FileExistsError: If path provided and file already exists at that location
        """
        doc = Document()

        if path is not None:
            abs_path = self._normalize_path(path)

            # Error-on-overwrite: check if file already exists
            if Path(abs_path).exists():
                raise FileExistsError(f"File already exists at {abs_path}")

            self._documents[abs_path] = doc
            return abs_path, doc
        else:
            # Generate temporary key
            key = self._next_untitled()
            self._documents[key] = doc
            return key, doc

    def open_document(self, path: str) -> Document:
        """
        Open an existing .docx file from disk into memory.

        If document is already open, returns the cached in-memory instance.

        Args:
            path: Path to .docx file to open

        Returns:
            Document object

        Raises:
            FileNotFoundError: If file doesn't exist at path
        """
        abs_path = self._normalize_path(path)

        # Return cached instance if already open
        if abs_path in self._documents:
            return self._documents[abs_path]

        # Check file exists before attempting to open
        if not Path(abs_path).exists():
            raise FileNotFoundError(f"File not found: {abs_path}")

        # Load from disk
        doc = Document(abs_path)
        self._documents[abs_path] = doc
        return doc

    def create_from_template(
        self,
        template_path: str,
        save_path: Optional[str] = None
    ) -> tuple[str, Document]:
        """
        Create a new document from a .docx or .dotx template.

        Args:
            template_path: Path to template file (.docx or .dotx)
            save_path: Optional path where document will be saved. If provided and
                       file exists, raises FileExistsError. If None, generates
                       temporary "Untitled-N" key.

        Returns:
            Tuple of (key, Document) where key is save_path or "Untitled-N"

        Raises:
            FileNotFoundError: If template doesn't exist
            FileExistsError: If save_path provided and file already exists
        """
        abs_template = self._normalize_path(template_path)

        # Check template exists
        if not Path(abs_template).exists():
            raise FileNotFoundError(f"Template not found: {abs_template}")

        # Create document from template
        doc = Document(abs_template)

        if save_path is not None:
            abs_save = self._normalize_path(save_path)

            # Error-on-overwrite
            if Path(abs_save).exists():
                raise FileExistsError(f"File already exists at {abs_save}")

            self._documents[abs_save] = doc
            return abs_save, doc
        else:
            # Generate temporary key
            key = self._next_untitled()
            self._documents[key] = doc
            return key, doc

    def save_document(self, path: str, save_as: Optional[str] = None):
        """
        Save an open document to disk.

        Args:
            path: Current key/path of open document
            save_as: Optional new path for save-as operation. If provided, document
                     is saved to new location and re-keyed in the documents dict.

        Raises:
            ValueError: If document is not currently open
        """
        # Normalize current path
        current_key = path if path.startswith("Untitled-") else self._normalize_path(path)

        if current_key not in self._documents:
            raise ValueError(f"Document not open: {path}")

        doc = self._documents[current_key]

        if save_as is not None:
            # Save-as: save to new path and re-key
            abs_new = self._normalize_path(save_as)

            # Create parent directories if needed
            Path(abs_new).parent.mkdir(parents=True, exist_ok=True)

            doc.save(abs_new)

            # Re-key in dictionary (remove old key, add new)
            del self._documents[current_key]
            self._documents[abs_new] = doc
        else:
            # Regular save: save to current path
            if current_key.startswith("Untitled-"):
                raise ValueError(
                    f"Cannot save untitled document without path. Use save_as parameter."
                )

            # Create parent directories if needed
            Path(current_key).parent.mkdir(parents=True, exist_ok=True)

            doc.save(current_key)

    def close_document(self, path: str):
        """
        Close an open document (remove from memory).

        Unsaved changes are discarded. This is explicit-only behavior.

        Args:
            path: Key/path of document to close

        Raises:
            ValueError: If document is not currently open
        """
        # Normalize path
        key = path if path.startswith("Untitled-") else self._normalize_path(path)

        if key not in self._documents:
            raise ValueError(f"Document not open: {path}")

        del self._documents[key]

    def get_document(self, path: str) -> Document:
        """
        Get an open document by its key/path.

        Args:
            path: Key/path of document

        Returns:
            Document object

        Raises:
            ValueError: If document is not currently open
        """
        # Normalize path
        key = path if path.startswith("Untitled-") else self._normalize_path(path)

        if key not in self._documents:
            raise ValueError(f"Document not open: {path}")

        return self._documents[key]

    def list_documents(self) -> list[str]:
        """
        List all currently open document keys/paths.

        Returns:
            List of document keys (absolute paths or "Untitled-N" strings)
        """
        return list(self._documents.keys())

    def close_all(self):
        """
        Close all open documents (cleanup for server shutdown).

        Clears all documents from memory and resets the untitled counter.
        This is called during graceful server shutdown to release resources.

        Note: Unsaved changes are discarded (consistent with close_document behavior).
        """
        doc_count = len(self._documents)
        if doc_count > 0:
            self._documents.clear()
            self._untitled_counter = 0
            return doc_count
        return 0


# Module-level singleton
document_manager = DocumentManager()
